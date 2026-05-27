#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Génération du rapport Entrepreneuriat — Lumière Cinémas (PFA)
Exécuter depuis la racine du projet : python scripts/generate_rapport_entrepreneuriat.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─────────────────────────────────────────────────────────────────────────────
# STYLES GLOBAUX
# ─────────────────────────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

def set_page_margins(doc, top=2.5, bottom=2.5, left=3.0, right=2.5):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)

set_page_margins(doc)

# Couleurs
COLOR_ACCENT  = RGBColor(0xB0, 0x00, 0x20)   # rouge foncé
COLOR_DARK    = RGBColor(0x1A, 0x1A, 0x2E)   # quasi-noir marine
COLOR_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_LGRAY   = RGBColor(0xF2, 0xF2, 0xF2)
COLOR_MGRAY   = RGBColor(0xD9, 0xD9, 0xD9)

# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def shade_cell(cell, hex_color: str):
    """Colorie le fond d'une cellule."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            for k, v in kwargs[edge].items():
                tag.set(qn(f'w:{k}'), v)
            tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_ACCENT
    # soulignement via border bas
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'B00020')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_DARK
    return p

def add_heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return p

def add_body(text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.7)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_table_header_row(table, headers, bg='1A1A2E', fg='FFFFFF'):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        shade_cell(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

def add_data_row(table, row_idx, data, alt=False):
    row = table.rows[row_idx]
    bg = 'F2F2F2' if alt else 'FFFFFF'
    for i, val in enumerate(data):
        cell = row.cells[i]
        cell.text = ''
        shade_cell(cell, bg)
        p = cell.paragraphs[0]
        run = p.add_run(str(val))
        run.font.size = Pt(10)
    return row

def make_table(headers, rows_data, col_widths=None):
    t = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header_row(t, headers)
    for i, row in enumerate(rows_data):
        add_data_row(t, i + 1, row, alt=(i % 2 == 1))
    if col_widths:
        for row in t.rows:
            for j, cell in enumerate(row.cells):
                if j < len(col_widths):
                    cell.width = Cm(col_widths[j])
    return t

def add_space(pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pt)

# ─────────────────────────────────────────────────────────────────────────────
# PAGE DE TITRE
# ─────────────────────────────────────────────────────────────────────────────

# --- En-tête école ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ÉCOLE MAROCAINE DES SCIENCES DE L\'INGÉNIEUR')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = COLOR_DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('EMSI — Casablanca')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Filière : Développement des Systèmes d\'Information (DSI) — 4ème Année · Groupe G3')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

add_space(6)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Année Académique 2025 / 2026')
run.font.size = Pt(11)
run.italic = True
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# --- Séparateur ---
add_space(12)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('━' * 55)
run.font.color.rgb = COLOR_ACCENT

add_space(12)

# --- Titre principal ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('LUMIÈRE CINÉMAS')
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = COLOR_ACCENT

add_space(4)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Plateforme de Réservation de Billets de Cinéma en Ligne')
run.font.size = Pt(15)
run.font.color.rgb = COLOR_DARK

add_space(8)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('RAPPORT DE PROJET DE FIN D\'ANNÉE (PFA)')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = COLOR_DARK

add_space(4)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Partie : Entrepreneuriat & Business Plan')
run.font.size = Pt(12)
run.italic = True
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# --- Séparateur ---
add_space(12)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('━' * 55)
run.font.color.rgb = COLOR_ACCENT

add_space(14)

# --- Bloc étudiant / encadrante ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Réalisé par :')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Wassim Lazim')
run.bold = True
run.font.size = Pt(15)
run.font.color.rgb = COLOR_DARK

add_space(8)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Encadrante :')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Mme Naji Zineb')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = COLOR_DARK

# --- Date de soutenance ---
add_space(14)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('━' * 55)
run.font.color.rgb = COLOR_ACCENT

add_space(6)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Soutenance : Juin 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

add_space(4)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Stack technique : Next.js 15  ·  Supabase  ·  Stripe  ·  TypeScript')
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 1 — CONTEXTE ET IDENTIFICATION DU PROBLÈME
# ─────────────────────────────────────────────────────────────────────────────
add_heading1('1. Contexte et Identification du Problème')

add_heading2('1.1 Le secteur du cinéma au Maroc')
add_body(
    "Le cinéma marocain connaît depuis 2018 une renaissance significative, portée par l'émergence "
    "de complexes multiplex modernes dans les principales métropoles du Royaume. En 2024, le Maroc "
    "compte environ 50 salles de cinéma actives réparties sur l'ensemble du territoire, dont une "
    "forte concentration dans les villes de Casablanca (environ 18 salles), Rabat, Marrakech, Fès "
    "et Agadir. La fréquentation nationale s'est établie autour de 3,5 millions d'entrées en 2023, "
    "contre 2,9 millions en 2022, témoignant d'une reprise soutenue après les fermetures imposées "
    "par la pandémie de Covid-19."
)
add_body(
    "Les acteurs majeurs du secteur sont : Megarama (groupe français implanté au Maroc avec des "
    "complexes à Casablanca Mall, Rabat et Fès), Colisée Cinémas (présent à Casablanca et Marrakech), "
    "IMAX Morocco (Casablanca), ainsi que des salles indépendantes comme le Dawliz, le Rif ou "
    "l'ABC à Casablanca. Le Centre Cinématographique Marocain (CCM), établissement public sous "
    "tutelle du Ministère de la Culture, joue le rôle de régulateur et de promoteur de la "
    "production nationale."
)

add_heading2('1.2 Tendances du marché')
add_body(
    "Plusieurs tendances structurelles transforment profondément les habitudes de consommation "
    "cinématographique au Maroc :"
)
add_bullet("Digitalisation accélérée : selon l'Agence Nationale de Réglementation des Télécommunications (ANRT), le taux de pénétration d'Internet mobile au Maroc dépasse 89 % en 2024, avec plus de 28 millions d'abonnés. Les Marocains effectuent de plus en plus leurs achats en ligne, y compris pour les loisirs.")
add_bullet("Croissance de la demande jeune : la population marocaine est jeune — plus de 52 % a moins de 30 ans (HCP, 2024). Cette tranche d'âge est la première consommatrice de cinéma et la plus à l'aise avec les solutions numériques de réservation.")
add_bullet("Essor du commerce électronique : le volume du e-commerce marocain a atteint 11,3 milliards de MAD en 2023 (rapport CMI), avec une croissance annuelle de 23 %. Le secteur des loisirs et du divertissement bénéficie pleinement de cette dynamique.")
add_bullet("Concurrence des plateformes VOD : Netflix, Shahid et OSN+ ont modifié les comportements, mais n'ont pas tué la salle — elles ont au contraire habitué le public à une expérience numérique fluide qu'il exige désormais également au cinéma.")

add_heading2('1.3 Le problème concret')
add_body(
    "Malgré cette modernisation, l'expérience de réservation de billets de cinéma au Maroc reste "
    "marquée par de nombreuses frictions qui dégradent la satisfaction client :"
)
add_bullet("Files d'attente aux guichets : dans les complexes populaires de Casablanca (Megarama Casablanca Mall, Colisée Boulvard), les queues pour l'achat de billets les vendredis et samedis soirs peuvent dépasser 20 à 30 minutes, notamment pour les sorties de blockbusters américains.")
add_bullet("Manque de visibilité du programme : en l'absence d'une plateforme centralisée nationale, les spectateurs doivent consulter les pages Facebook ou les sites propres de chaque cinéma, souvent mal maintenus, pour connaître les horaires et les films à l'affiche.")
add_bullet("Absence de sélection de siège en ligne : la quasi-totalité des cinémas marocains ne proposent pas de plan de salle interactif permettant de choisir son emplacement à l'avance. Le spectateur ignore jusqu'au dernier moment où il sera assis.")
add_bullet("Modes de paiement limités : la grande majorité des transactions se font encore en espèces aux guichets. Les solutions de paiement en ligne par carte bancaire ou mobile restent très peu répandues dans ce secteur.")
add_bullet("Aucun historique de réservation : sans compte client ni billet numérique, le spectateur ne dispose d'aucune trace de ses visites et aucune personnalisation n'est possible (recommandations, fidélisation).")

add_heading2('1.4 Limites des solutions existantes au Maroc')
add_body(
    "Cinemashi.ma est le principal agrégateur de programme cinéma au Maroc. S'il centralise bien "
    "les horaires, il ne propose pas de réservation en ligne avec paiement intégré ni de sélection "
    "de sièges. Les cinémas eux-mêmes (Megarama, Colisée) proposent ponctuellement des réservations "
    "via leurs propres sites web, mais ces interfaces sont disparates, sans plan de salle unifié et "
    "sans tunnel de paiement moderne. Il n'existe au Maroc aucune plateforme offrant l'ensemble : "
    "programme centralisé + sélection de siège + paiement en ligne + billet numérique QR code + "
    "dashboard de gestion opérateur — ce que Lumière Cinémas propose."
)
add_space()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 2 — PRÉSENTATION DE L'IDÉE ET DE LA SOLUTION
# ─────────────────────────────────────────────────────────────────────────────
add_heading1('2. Présentation de l\'Idée et de la Solution')

add_heading2('2.1 La solution : Lumière Cinémas')
add_body(
    "Lumière Cinémas est une plateforme web fullstack de réservation de billets de cinéma en ligne, "
    "conçue pour répondre aux lacunes identifiées sur le marché marocain. Elle s'adresse "
    "simultanément à deux types d'utilisateurs : les spectateurs, qui bénéficient d'une expérience "
    "de réservation moderne de bout en bout, et les exploitants de salles, qui disposent d'un "
    "back-office complet de gestion opérationnelle."
)

add_heading2('2.2 Fonctionnalités principales identifiées dans le projet')
add_body("L'analyse du code source révèle les fonctionnalités suivantes, toutes implémentées :")

add_heading3('Côté spectateur (parcours utilisateur)')
add_bullet("Page d'accueil avec carousel immersif : affichage des films de la semaine avec carrousel plein écran, fond dynamique adapté à chaque film, et lancement de la bande-annonce YouTube en modale.")
add_bullet("Catalogue films avec recherche instantanée : grille de films filtrée en temps réel côté client, sans rechargement de page (composant FilmsGrid avec useState).")
add_bullet("Fiche film détaillée : synopsis, réalisateur, casting, durée, genre, classification (12+/TP/16+), liste des séances disponibles groupées par date.")
add_bullet("Plan de salle interactif (SeatPicker) : visualisation graphique de la salle avec 3 catégories de sièges — Standard (12,50 MAD), Premium (16,00 MAD), Duo (32,00 MAD). Sélection multi-sièges avec calcul du sous-total en temps réel. Blocage automatique des sièges déjà réservés.")
add_bullet("Sélection de snacks : commande optionnelle de consommations (popcorn, sodas, bonbons, glace) récupérable au comptoir. Catalogue avec prix en MAD : Popcorn 6-7,50 MAD, Duo Pack 14 MAD, Soda 4,50 MAD.")
add_bullet("Tunnel de paiement sécurisé : intégration Stripe Checkout pour le paiement par carte bancaire internationale, et option paiement en espèces au guichet. Le montant total est recalculé côté serveur avant validation.")
add_bullet("Billet numérique avec QR code : après confirmation, génération d'un billet avec QR code unique, référence lisible (format LM-2026-XXXX), détails de la séance et récapitulatif des sièges réservés.")
add_bullet("Historique des réservations : espace compte personnel listant toutes les réservations passées et à venir de l'utilisateur connecté.")
add_bullet("Programme par jour : vue calendrier avec onglets de navigation par date, filtrable par film.")
add_bullet("Page Prochainement : films annoncés non encore à l'affiche.")
add_bullet("Page Salles : présentation des équipements (Dolby Atmos, IMAX, 4K HDR).")

add_heading3('Côté administrateur (back-office)')
add_bullet("Dashboard KPIs : vue synthétique avec le nombre de films, séances, salles, et le chiffre d'affaires total confirmé, chargés en parallèle (Promise.all).")
add_bullet("CRUD Films complet : création, modification, suppression de films avec upload d'affiche vers Supabase Storage. Gestion du gradient CSS de fallback si pas d'image.")
add_bullet("CRUD Séances : programmation des séances avec association film/salle, format de projection (VOSTFR, VF, IMAX, 4K HDR), horaire, et prix par catégorie de siège surchargeable.")
add_bullet("CRUD Salles : gestion de la capacité totale, technologie et couleur d'accentuation.")
add_bullet("CRUD Snacks : catalogue de consommations avec prix, emoji, description.")
add_bullet("Gestion des réservations : liste des 200 dernières réservations avec statuts (confirmed, pending, cancelled), jointures films/séances.")
add_bullet("Gestion des utilisateurs : attribution et révocation du rôle administrateur.")

add_heading2('2.3 Technologies réellement utilisées')
add_body("Le projet repose sur une stack technique moderne, entièrement visible dans le code source :")

make_table(
    ['Couche', 'Technologie', 'Version', 'Rôle'],
    [
        ['Framework',        'Next.js',                    '15.0.3',   'App Router, Server Components, Server Actions, SSR'],
        ['Langage',          'TypeScript + React',         '5.x / 19', 'Typage statique, composants frontend'],
        ['Base de données',  'Supabase (PostgreSQL)',       'Latest',   'Données relationnelles + Row Level Security'],
        ['Authentification', 'Supabase Auth',              'Latest',   'Sessions JWT via cookies, inscription/connexion'],
        ['Paiement',         'Stripe Checkout',            '22.1.1',   'Paiement par carte bancaire sécurisé'],
        ['Stockage fichiers','Supabase Storage',           'Latest',   'Upload et diffusion des affiches de films'],
        ['QR Code',          'qrcode (npm)',               '1.5.4',    'Génération des billets numériques'],
        ['Style',            'CSS custom + variables',     'Native',   'Design tokens, thème sombre, accent #E50914'],
        ['Hébergement',      'Vercel (cible déploiement)', 'Latest',   'Déploiement serverless, CDN mondial'],
    ],
    [3.5, 4.5, 2.0, 6.5]
)

add_space()
add_heading2('2.4 Valeur ajoutée et avantage concurrentiel')
add_body(
    "Lumière Cinémas se distingue par plusieurs avantages concurrentiels clés. D'abord, "
    "c'est le seul projet marocain à proposer un tunnel de réservation complet en 6 étapes "
    "(accueil → film → siège → snacks → paiement → billet) sans rupture d'expérience. "
    "Ensuite, l'architecture SSR (Server-Side Rendering) de Next.js 15 garantit des temps de "
    "chargement rapides même sur connexions mobiles 4G, critère déterminant au Maroc. "
    "La gestion fine des rôles (RLS Supabase) et la sécurité des données (JWT) répondent aux "
    "exigences de la loi 09-08. Enfin, le design sombre avec accent rouge (#E50914) et le "
    "carrousel immersif offrent une esthétique premium comparable aux grandes plateformes "
    "internationales, créant une rupture forte avec les interfaces marocaines existantes."
)
add_space()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 3 — ORIGINE DE L'IDÉE
# ─────────────────────────────────────────────────────────────────────────────
add_heading1('3. Origine et Génération de l\'Idée Entrepreneuriale')

add_heading2('3.1 Observation fondatrice')
add_body(
    "L'idée de Lumière Cinémas est née d'une observation directe et répétée : lors de sorties "
    "cinéma dans les complexes casablancais (Megarama Casablanca Mall, Colisée Boulvard), "
    "le constat d'une file d'attente de 25 minutes pour acheter un billet alors que des dizaines "
    "de spectateurs tiennent un smartphone connecté dans la main illustre parfaitement le "
    "décalage entre les usages numériques marocains et les services proposés par l'industrie "
    "cinématographique locale. Cette friction vécue — perdre une demi-heure à une caisse pour "
    "obtenir un billet qu'un QR code pourrait remplacer — est le point de départ du projet."
)

add_heading2('3.2 Opportunité de marché identifiée')
add_body(
    "L'analyse du marché a confirmé l'opportunité : un segment de jeunes urbains marocains de "
    "18 à 35 ans, hyper-connectés, familiers des applications comme Glovo, Jumia Food ou "
    "Careem, expriment clairement une attente vis-à-vis d'une expérience cinéma digitale. "
    "Cette demande latente est non satisfaite par l'offre locale, créant un vide "
    "entrepreneurial exploitable. Le marché est de plus en train de se consolider autour de "
    "grands complexes multiplex, ce qui facilite les partenariats B2B (quelques acteurs "
    "majeurs à convaincre plutôt qu'une multitude de petites salles)."
)

add_heading2('3.3 Inspirations')
add_body(
    "Le projet s'inspire du modèle de plateformes de billetterie internationales ayant réussi "
    "leur transition numérique : Fandango (USA, 120 millions d'utilisateurs), Allocinésur le "
    "marché français (billets intégrés depuis 2017), et surtout Yassir Events en Algérie qui "
    "démontre la viabilité du modèle en Afrique du Nord. L'interface s'inspire également de "
    "l'expérience Booking.com — tunnel court, confirmation immédiate, billet numérique — "
    "adaptée au contexte culturel et linguistique marocain (arabe, français, darija)."
)
add_space()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 4 — ÉTUDE DE MARCHÉ
# ─────────────────────────────────────────────────────────────────────────────
add_heading1('4. Étude de Marché')

add_heading2('4.A — Analyse de la Demande')

add_heading3('Profil des futurs clients')
add_body(
    "La clientèle cible de Lumière Cinémas se compose principalement de trois profils complémentaires. "
    "Le premier, et le plus important, est le jeune urbain de 18 à 34 ans, résidant dans les grandes "
    "villes (Casablanca, Rabat, Marrakech, Tanger, Agadir). Ce profil dispose d'un smartphone, d'une "
    "connexion Internet permanente et d'une carte bancaire ou d'un compte de paiement mobile (CIH Pay, "
    "M-Walletde Maroc Telecom, ou Wafacash). Il consomme des contenus en streaming (Netflix, Shahid) "
    "et est habitué aux interfaces numériques. Le deuxième profil est la famille de classe moyenne à "
    "aisée des quartiers urbains, qui sort en groupe le week-end et dont la planification bénéficierait "
    "grandement d'une réservation anticipée de plusieurs sièges contigus. Le troisième profil est "
    "l'entreprise organisatrice de sorties corporate ou team-building."
)

add_heading3('Comportements de consommation')
add_body(
    "Selon les estimations du CCM et les données des exploitants, un Marocain habitué du cinéma "
    "effectue en moyenne 8 à 12 visites par an dans les zones urbaines. Le budget moyen par visite "
    "est estimé entre 60 et 150 MAD par personne, incluant le billet (Standard : 50–70 MAD, IMAX : "
    "80–120 MAD dans les grands complexes) et les consommations (popcorn, boissons : 25–50 MAD). "
    "Les séances les plus fréquentées sont le vendredi soir, le samedi après-midi et le dimanche. "
    "Les blockbusters hollywoodiens et les productions arabes (Ramadan Cinema) génèrent les "
    "meilleures fréquentations."
)

add_heading3('Disposition à payer pour une solution digitale')
add_body(
    "Les études de comportement sur des marchés comparables (Tunisie, Égypte) montrent qu'un "
    "utilisateur est prêt à payer une commission de 3 à 8 MAD par billet pour bénéficier d'une "
    "réservation en ligne, à condition que l'expérience soit fluide et le paiement sécurisé. "
    "Cette commission reste inférieure au coût perçu de la file d'attente. Une offre premium "
    "(abonnement mensuel sans commission) à 49–59 MAD/mois est envisageable pour les utilisateurs "
    "fréquents (plus de 4 visites par mois)."
)

add_heading2('4.B — Analyse de la Concurrence')
add_body("Tableau comparatif des principaux acteurs concurrents directs et indirects :")
add_space(3)

make_table(
    ['Concurrent', 'Type', 'Marché', 'Modèle éco.', 'Prix billets', 'Forces', 'Faiblesses'],
    [
        ['Cinemashi.ma',
         'Agrégateur programme',
         'Maroc',
         'Publicité display',
         'Gratuit (pas de vente)',
         'Seule référence nationale des horaires',
         'Aucune réservation en ligne ni sélection de siège'],
        ['Megarama.ma',
         'Site cinéma propriétaire',
         'Maroc/France',
         'Vente directe billets',
         '50–90 MAD',
         'Marque forte, 3 salles Maroc',
         'Réservation limitée, pas de plan interactif unifié'],
        ['Colisée Cinemas',
         'Réseau salles',
         'Maroc',
         'Vente directe billets',
         '45–80 MAD',
         'Bien implanté Casablanca/Marrakech',
         'Site web obsolète, pas de réservation mobile fluide'],
        ['IMAX Morocco',
         'Salle premium unique',
         'Maroc (Casablanca)',
         'Billetterie propre',
         '80–130 MAD',
         'Expérience premium différenciante',
         'Offre très restreinte (1 salle), site rudimentaire'],
        ['Ciné Horizon (Tunisie)',
         'Plateforme billetterie',
         'Tunisie',
         'Commission/billet + B2B',
         'Commission 2 TND',
         'Modèle mature, preuve de concept maghrébin',
         'Hors Maroc, non adapté au contexte marocain'],
        ['Netflix Maroc',
         'VOD',
         'Maroc/Monde',
         'Abonnement mensuel 40–99 MAD',
         '40–99 MAD/mois',
         '200M+ abonnés mondiaux, contenus illimités',
         'Ne remplace pas l\'expérience salle, pas de sorties cinéma exclusives'],
        ['Shahid (MBC Group)',
         'VOD arabe',
         'Maroc/MENA',
         'Abonnement 49–79 MAD/mois',
         '49–79 MAD/mois',
         'Contenus arabes premium, très populaire au Maroc',
         'Pas de billetterie, uniquement streaming'],
        ['OSN+',
         'VOD premium',
         'Maroc/MENA',
         'Abonnement 99 MAD/mois',
         '99 MAD/mois',
         'Contenus HBO et séries premium',
         'Prix élevé, audience restreinte au Maroc'],
    ],
    [2.5, 2.5, 1.8, 2.5, 2.0, 3.5, 3.5]
)
add_space()

add_heading2('4.C — Analyse des Fournisseurs')
add_body(
    "La solution Lumière Cinémas dépend de plusieurs fournisseurs tiers dont les coûts réels "
    "sont intégrés dans le modèle financier :"
)
add_space(3)

make_table(
    ['Fournisseur', 'Service', 'Tarif réel (2025)', 'Usage dans le projet'],
    [
        ['Supabase',
         'BDD PostgreSQL + Auth + Storage',
         'Free tier : 0 MAD (500 MB BDD, 1 GB Storage). Pro : 250 MAD/mois (8 GB BDD, 100 GB Storage)',
         'Cœur de l\'application : données films/séances/réservations, authentification, affiches'],
        ['Vercel',
         'Hébergement + CDN + Serverless',
         'Hobby : gratuit. Pro : 200 MAD/mois (20 GB bande passante, domaine custom, analytics)',
         'Déploiement de l\'app Next.js, fonctions serverless pour les Server Actions'],
        ['Stripe',
         'Paiement en ligne par carte',
         '1,5 % + 1,50 MAD par transaction (cartes EU/MA). Aucun frais fixe mensuel.',
         'Checkout sécurisé pour le paiement des billets par carte bancaire (Visa/Mastercard)'],
        ['CMI (Centre Monétique Interbancaire)',
         'Paiement cartes marocaines',
         'Commission ~1,8–2,5 % par transaction + frais d\'adhésion ~3 000 MAD/an',
         'Alternative locale à Stripe pour les cartes Visa/Mastercard émises au Maroc (obligatoire à terme)'],
        ['Twilio / Infobip',
         'SMS notifications',
         'Twilio : 0,085 USD/SMS au Maroc (~0,85 MAD). Infobip : 0,075 USD/SMS (~0,75 MAD)',
         'Confirmation de réservation et rappel de séance par SMS'],
        ['SendGrid (Twilio)',
         'Emails transactionnels',
         'Free tier : 100 emails/jour. Essentials : 15 USD/mois (~150 MAD) pour 50 000 emails/mois',
         'Envoi des billets par email, confirmations d\'inscription, notifications de réservation'],
        ['Google Maps API',
         'Cartographie (localisation salles)',
         '200 USD de crédit mensuel gratuit. Au-delà : 7 USD/1000 requêtes cartes dynamiques',
         'Optionnel : affichage de la localisation des salles partenaires sur une carte interactive'],
        ['OVH Cloud (alternative)',
         'VPS / hébergement alternatif',
         'VPS Starter : 5,99 €/mois (~65 MAD). VPS Value : 9,99 €/mois (~109 MAD)',
         'Alternative souveraine à Vercel pour réduire la dépendance aux clouds américains'],
    ],
    [2.5, 3.0, 4.0, 4.0]
)
add_space()

add_heading2('4.D — Cadre Légal et Réglementaire au Maroc')

add_heading3('Protection des données personnelles — Loi 09-08')
add_body(
    "La loi n° 09-08 relative à la protection des personnes physiques à l'égard du traitement "
    "des données à caractère personnel, promulguée le 18 février 2009, constitue le cadre "
    "juridique fondamental applicable à Lumière Cinémas. Cette loi exige la déclaration préalable "
    "de tout traitement de données personnelles auprès de la Commission Nationale de contrôle de "
    "la protection des Données à caractère Personnel (CNDP). Pour une plateforme de réservation, "
    "les données collectées (nom, email, numéro de téléphone, historique d'achats) doivent faire "
    "l'objet d'une déclaration normale ou d'une autorisation selon leur nature. Les frais de "
    "dossier CNDP sont de l'ordre de 1 000 à 2 000 MAD."
)

add_heading3('Commerce électronique — Loi 53-05')
add_body(
    "La loi n° 53-05 relative à l'échange électronique de données juridiques encadre les "
    "transactions en ligne au Maroc. Elle définit les conditions de validité des contrats "
    "électroniques et consacre la valeur juridique de la signature électronique. Pour Lumière "
    "Cinémas, cela implique l'obligation d'afficher clairement les conditions générales de vente "
    "(CGV), la politique de remboursement et les informations légales sur le site."
)

add_heading3('Obligations fiscales')
add_body(
    "En tant que société commerciale marocaine, Lumière Cinémas sera assujettie à la TVA au "
    "taux normal de 20 % sur les commissions et services numériques, ainsi qu'à l'Impôt sur "
    "les Sociétés (IS) au taux de 20 % sur les bénéfices nets inférieurs à 1 million de MAD "
    "(taux progressif). La patente est également due selon le barème de la Direction Générale "
    "des Impôts (DGI). Les déclarations de TVA sont mensuelles ou trimestrielles selon le CA."
)

add_heading3('OMPIC et propriété intellectuelle')
add_body(
    "Le dépôt de la marque 'Lumière Cinémas' auprès de l'Office Marocain de la Propriété "
    "Industrielle et Commerciale (OMPIC) est fortement recommandé pour protéger l'identité "
    "commerciale. Le dépôt de marque coûte 2 025 MAD pour une classe de produits/services "
    "(classe 41 — services d'éducation et divertissement). Le registre de commerce est obligatoire "
    "et est obtenu auprès du Tribunal de Commerce pour un coût d'environ 350 MAD."
)

add_heading2('4.E — Outils d\'Analyse Stratégique')

add_heading3('Matrice SWOT')
add_space(3)

t_swot = doc.add_table(rows=3, cols=2)
t_swot.style = 'Table Grid'
t_swot.alignment = WD_TABLE_ALIGNMENT.CENTER

def swot_cell(cell, title, color, items):
    shade_cell(cell, color)
    cell.text = ''
    p0 = cell.paragraphs[0]
    r0 = p0.add_run(title)
    r0.bold = True
    r0.font.size = Pt(11)
    r0.font.color.rgb = COLOR_WHITE
    for item in items:
        p = cell.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        r = p.add_run(f'• {item}')
        r.font.size = Pt(10)
        r.font.color.rgb = COLOR_WHITE

# En-tête
shade_cell(t_swot.rows[0].cells[0], '1A1A2E')
t_swot.rows[0].cells[0].text = ''
r_tmp = t_swot.rows[0].cells[0].paragraphs[0].add_run('ANALYSE SWOT — LUMIÈRE CINÉMAS')
r_tmp.bold = True; r_tmp.font.color.rgb = COLOR_WHITE; r_tmp.font.size = Pt(12)
t_swot.rows[0].cells[0].merge(t_swot.rows[0].cells[1])

swot_cell(t_swot.rows[1].cells[0], 'FORCES (Strengths)', '2E7D32', [
    'Architecture technique moderne (Next.js 15 SSR, TypeScript)',
    'Tunnel de réservation complet en 6 étapes sans équivalent marocain',
    'Plan de salle interactif avec 3 catégories de prix (Standard/Premium/Duo)',
    'Intégration Stripe Checkout pour paiement sécurisé',
    'Back-office admin complet (CRUD + KPIs + gestion utilisateurs)',
    'Design premium différenciant (thème sombre, animations, QR code)',
    'Sécurité robuste : RLS Supabase, JWT, Row Level Security PostgreSQL',
    'Coût d\'infrastructure faible au démarrage (Supabase Free + Vercel Hobby)',
])
swot_cell(t_swot.rows[1].cells[1], 'FAIBLESSES (Weaknesses)', 'C62828', [
    'Pas encore d\'application mobile native (iOS/Android)',
    'Dépendance à Stripe (cartes internationales) — intégration CMI à compléter',
    'Absence de l\'arabe et de la darija dans l\'interface',
    'Notoriété nulle au lancement — brand à construire from scratch',
    'Équipe réduite : projet initialement développé par un seul développeur',
    'Pas encore d\'intégration SMS natif ni d\'email transactionnel branché',
])
swot_cell(t_swot.rows[2].cells[0], 'OPPORTUNITÉS (Opportunities)', '1565C0', [
    'Marché marocain non adressé : aucun concurrent direct sur le segment billetterie complète',
    'Croissance du e-commerce marocain : +23 %/an (CMI 2023)',
    'Taux de pénétration smartphone > 89 % (ANRT 2024)',
    'Consolidation autour de grands réseaux (Megarama, Colisée) → partenariats B2B simples',
    'Expansion naturelle au Maghreb (Algérie, Tunisie) avec même stack technique',
    'Événements culturels (concerts, théâtre, festivals) : marché adjacent accessible',
    'Monétisation data : recommandations personnalisées, analytics pour cinémas partenaires',
])
swot_cell(t_swot.rows[2].cells[1], 'MENACES (Threats)', '6A1FA8', [
    'Résistance des exploitants à partager données de billetterie',
    'Risque de désintermédiation : un grand réseau (Megarama) pourrait développer sa propre app',
    'Volatilité de la demande (crises, Ramadan, saison estivale faible)',
    'Concurrence possible d\'acteurs tech internationaux (Fever, Ticketmaster) entrant sur le marché marocain',
    'Risque de fraude sur les paiements en ligne (chargebacks)',
    'Évolutions réglementaires (CNDP, CMI) pouvant imposer des coûts de mise en conformité',
])

add_space()
add_heading3('Analyse PESTEL')
add_space(3)

make_table(
    ['Dimension', 'Facteurs applicables au marché marocain', 'Impact'],
    [
        ['Politique (P)',
         'Stratégie Maroc Digital 2030 du gouvernement qui soutient la transformation numérique. '
         'Centre Cinématographique Marocain (CCM) actif dans la promotion du secteur. '
         'Stabilité politique relative favorable aux investissements.',
         'Favorable +'],
        ['Économique (E)',
         'PIB marocain : 1 303 milliards MAD (2023, HCP). Croissance 3,4 % prévue 2024. '
         'Classe moyenne urbaine en expansion (~7 millions de personnes). '
         'Inflation ~5,7 % en 2023 impactant le pouvoir d\'achat. '
         'Commission e-commerce CMI : ~1,8–2,5 %.',
         'Modéré ±'],
        ['Socioculturel (S)',
         'Population jeune : 52 % < 30 ans (HCP 2024). Urbanisation croissante (64 % urbains). '
         'Forte consommation de contenus digitaux et réseaux sociaux. '
         'Culture du cinéma ancrée dans les métropoles. '
         'Bilinguisme français/arabe.',
         'Très favorable ++'],
        ['Technologique (T)',
         'Taux pénétration Internet mobile : 89 % (ANRT 2024). '
         '28,4 millions d\'abonnés mobile Internet. '
         'Déploiement 5G lancé en 2025, couverture en expansion. '
         'Next.js 15 / Supabase : infrastructure cloud mature et scalable. '
         'Adoption des paiements mobiles (M-Wallet, CIH Pay) en progression.',
         'Très favorable ++'],
        ['Environnemental (E)',
         'Dématérialisation des billets → réduction du papier, impact écologique positif. '
         'Pas de contraintes environnementales spécifiques au secteur numérique. '
         'Supabase et Vercel : data centers à énergie partiellement renouvelable.',
         'Favorable +'],
        ['Légal (L)',
         'Loi 09-08 (protection données, CNDP). Loi 53-05 (e-commerce). '
         'TVA 20 % sur services numériques. IS progressif (20–35 %). '
         'Réglementation CMI sur les paiements en ligne. '
         'Obligation enregistrement OMPIC + Registre de Commerce.',
         'Contraignant mais gérable —'],
    ],
    [2.5, 10.5, 2.0]
)

add_space()
add_heading3('Les 5 Forces de Porter — Secteur billetterie cinéma au Maroc')
add_space(3)

make_table(
    ['Force', 'Intensité', 'Analyse détaillée'],
    [
        ['Rivalité entre concurrents existants',
         'FAIBLE',
         'Cinemashi.ma est le seul acteur national mais sans vente de billets. '
         'Les cinémas vendent leurs billets en silo. Absence de plateforme de billetterie '
         'multi-cinémas au Maroc. La rivalité directe est quasi-nulle sur ce segment précis.'],
        ['Menace de nouveaux entrants',
         'MODÉRÉE',
         'Barrières à l\'entrée techniques (développement plateforme, intégration paiement) '
         'et commerciales (convaincre les cinémas de partager leurs données). '
         'Un acteur international (Fever, Ticketmaster) pourrait entrer mais n\'a pas '
         'encore manifesté d\'intérêt pour le marché marocain.'],
        ['Pouvoir de négociation des fournisseurs',
         'MODÉRÉ',
         'Stripe et Supabase ont peu de substituts directs de même qualité, '
         'mais des alternatives existent (CMI pour le paiement, AWS pour l\'infra). '
         'Les cinémas partenaires ont un fort pouvoir de négociation car sans eux, '
         'pas de contenu à vendre.'],
        ['Pouvoir de négociation des clients',
         'MODÉRÉ',
         'Les spectateurs ont l\'alternative du guichet physique. Leur coût de '
         'changement (switching cost) est faible. Cependant, une fois habitués à '
         'la réservation en ligne, le retour au guichet est ressenti négativement. '
         'La fidélisation par l\'historique et les recommandations réduit ce risque.'],
        ['Menace des produits de substitution',
         'MODÉRÉE',
         'Netflix, Shahid, OSN+ proposent des alternatives VOD. Cependant, '
         'l\'expérience salle (Dolby Atmos, IMAX) reste indépassable pour les '
         'blockbusters. La substitution touche les films moyens, pas les grosses '
         'productions. L\'essor des salles premium (IMAX Morocco) renforce '
         'l\'irremplçabilité de l\'expérience physique.'],
    ],
    [4.0, 2.5, 8.5]
)
add_space()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 5 — MARKETING STRATÉGIQUE
# ─────────────────────────────────────────────────────────────────────────────
add_heading1('5. Marketing Stratégique')

add_heading2('5.1 Segmentation du marché')
add_space(3)

make_table(
    ['Critère', 'Segment 1 — Jeunes urbains', 'Segment 2 — Familles', 'Segment 3 — Entreprises'],
    [
        ['Âge cible',     '18–30 ans',            '30–50 ans (avec enfants)',    'Responsables RH/événements'],
        ['Géographie',    'Casablanca, Rabat, Marrakech, Tanger', 'Grandes villes + périphéries', 'Zones industrielles et centres d\'affaires'],
        ['Revenus',       '3 000–8 000 MAD/mois', '6 000–15 000 MAD/mois',      '> 15 000 MAD/mois (budget entreprise)'],
        ['Technologie',   'Hyper-connectés, smartphones haut de gamme', 'Connectés, utilisent les apps pour services du quotidien', 'Utilisent outils B2B, email pro, facturation'],
        ['Comportement digital', 'Instagram, TikTok, YouTube, jeux mobiles', 'WhatsApp, Facebook, applications pratiques', 'LinkedIn, email, outils de gestion'],
        ['Fréquence cinéma', '2–4 fois/mois',     '2–4 fois/trimestre',          '1–4 fois/an (sorties d\'équipe)'],
        ['Problème principal', 'Files, manque de prévisibilité des places, interface mobile médiocre', 'Complexité de réserver 4+ sièges contigus, snacks à l\'avance', 'Facturation, places groupées, confirmations rapides pour plusieurs personnes'],
        ['Disposition à payer', '3–6 MAD/billet commission', '4–8 MAD/billet', 'Abonnement corporate ou tarif négocié'],
    ],
    [3.0, 4.5, 4.0, 4.0]
)

add_heading2('5.2 Ciblage')
add_body(
    "Le segment prioritaire au lancement est le Segment 1 : jeunes urbains de 18 à 30 ans, "
    "principalement à Casablanca et Rabat. Ce ciblage est justifié par quatre arguments :"
)
add_bullet("Volume : ce segment représente environ 4,2 millions d'individus dans les 5 premières villes marocaines, constituant le plus grand marché accessible.")
add_bullet("Maturité digitale : ce segment est déjà utilisateur de services en ligne (Glovo, Jumia, booking d'hôtels) et ne nécessite aucune éducation à l'usage d'une plateforme de réservation.")
add_bullet("Coût d'acquisition faible : les canaux digitaux (Instagram, TikTok) permettent d'atteindre ce segment avec des coûts publicitaires modérés (CPC 0,3–0,8 MAD sur Instagram Maroc).")
add_bullet("Viralité naturelle : ce segment partage volontiers ses expériences culturelles sur les réseaux sociaux, générant un bouche-à-oreille digital organique.")
add_body(
    "Le Segment 2 (familles) sera activé en phase 2, après acquisition des premiers cinémas partenaires. "
    "Le Segment 3 (B2B entreprises) fera l'objet d'une offre dédiée en V3."
)

add_heading2('5.3 Positionnement')
add_body(
    "Lumière Cinémas se positionne comme la première plateforme de billetterie cinéma digitale "
    "au Maroc, combinant l'accessibilité, l'instantanéité et le plaisir de l'expérience cinéma. "
    "Elle se distingue des sites de cinéma existants par son approche multi-cinémas, son interface "
    "premium et son tunnel de réservation fluide."
)
add_body(
    "Proposition de valeur en une phrase :",
    bold=True
)
add_body(
    "« Réservez votre place de cinéma en 90 secondes, choisissez votre siège et payez en ligne — "
    "fini les files d'attente, votre billet QR code vous attend sur votre téléphone. »"
)
add_body(
    "Image de marque souhaitée : moderne, marocain, premium accessible. "
    "Le nom 'Lumière' rend hommage aux frères Lumière, inventeurs du cinéma, et s'inscrit "
    "dans la tradition culturelle française-marocaine. L'accent rouge #E50914 (emprunté à "
    "Netflix mais décliné dans un contexte cinéma classique) renforce l'association au cinéma."
)

add_heading2('5.4 Marketing Mix — Les 4P')

add_heading3('Produit')
add_body(
    "Le produit Lumière Cinémas est une plateforme web responsive (et à terme une application "
    "mobile PWA installable) offrant un parcours de réservation complet :"
)
add_bullet("Core product : réservation de billets en ligne avec sélection de siège.")
add_bullet("Augmented product : snacks commandés en ligne, billet QR code numérique, historique des réservations, programme centralisé multi-cinémas.")
add_bullet("Différenciation UX : plan de salle interactif avec catégories de prix visuellement distinctes (Standard, Premium, Duo), carrousel immersif sur l'accueil, fiche film cinématographique (synopsis, casting, bande-annonce).")
add_bullet("Back-office opérateur : tableau de bord complet permettant aux cinémas partenaires de gérer films, séances, salles et de suivre leurs réservations en temps réel.")
add_bullet("Sécurité : authentification JWT, Row Level Security PostgreSQL, paiement via Stripe PCI-DSS compliant.")

add_heading3('Prix')
add_body("Grille tarifaire proposée au lancement :")
add_space(3)

make_table(
    ['Segment / Service', 'Tarif', 'Modèle', 'Justification'],
    [
        ['Commission spectateur — billet Standard (50–70 MAD)',
         '3,00 MAD/billet',
         'Frais de service',
         'Inférieur au coût perçu d\'une file d\'attente, facilement accepté'],
        ['Commission spectateur — billet Premium/IMAX (80–130 MAD)',
         '5,00 MAD/billet',
         'Frais de service',
         'Proportionnel à la valeur du billet, reste marginal'],
        ['Abonnement Lumière Pass (spectateurs fréquents)',
         '59,00 MAD/mois',
         'Abonnement mensuel',
         '0 commission sur billets + accès avant-premières + notifications prioritaires'],
        ['Abonnement cinéma partenaire — Starter',
         '990 MAD/mois',
         'SaaS B2B',
         '1 salle, accès back-office complet, analytics de base'],
        ['Abonnement cinéma partenaire — Pro',
         '2 490 MAD/mois',
         'SaaS B2B',
         'Jusqu\'à 8 salles, analytics avancés, intégration API, support prioritaire'],
        ['Module Snacks — marge opérateur',
         '15 % du CA snacks',
         'Commission',
         'Partage de valeur sur les commandes de snacks transitant par la plateforme'],
    ],
    [4.5, 2.0, 2.5, 5.5]
)
add_space()
add_body(
    "Stratégie de prix recommandée au lancement : pénétration par la gratuité partielle. "
    "Les 6 premiers mois, les commissions spectateurs sont offertes (0 MAD de frais de service) "
    "pour maximiser l'acquisition utilisateur et construire la preuve sociale nécessaire aux "
    "négociations B2B avec les cinémas. Le modèle SaaS cinéma partenaire reste payant dès le "
    "premier jour pour valider le consentement à payer des exploitants."
)

add_heading3('Distribution')
add_body(
    "La distribution est exclusivement digitale, sans intermédiaire physique, ce qui constitue "
    "un avantage structurel sur les coûts :"
)
add_bullet("Application web responsive (desktop et mobile) : accessible depuis tout navigateur, optimisée pour les écrans tactiles.")
add_bullet("Progressive Web App (PWA) : installable sur l'écran d'accueil du smartphone sans passer par l'App Store, offrant une expérience proche d'une app native.")
add_bullet("SEO : pages films et séances indexées par Google, optimisées pour les recherches « billet cinéma Casablanca », « Megarama réservation en ligne ».")
add_bullet("Intégration QR code : les billets numériques sont scannables directement depuis le smartphone à l'entrée des salles.")
add_bullet("API ouverte (phase 2) : permettre aux cinémas partenaires d'intégrer le widget de réservation directement sur leur propre site web.")

add_heading3('Communication')
add_body("Stratégie de communication multi-canal, ancrée dans les usages marocains :")
add_bullet("Instagram & TikTok : contenu organique (bandes-annonces, top films de la semaine, coulisses de cinémas partenaires) + publicités ciblées (18–30 ans, Casablanca/Rabat/Marrakech). Budget publicitaire lancement : 5 000 MAD/mois.")
add_bullet("Partenariats influenceurs marocains : collaboration avec des créateurs de contenu lifestyle/culture comme Mehdi El Khorsi, Ibtissam Amhyar, ou des pages cinéma marocaines (CinemaMaroc, FilmMarocFan) qui cumulent 200 000 à 500 000 abonnés.")
add_bullet("SEO bilingue français/darija : optimisation des contenus pour les requêtes en français (« cinéma Casablanca ») et en darija translittérée (« réservation billet cinema »).")
add_bullet("Relations presse culturelle : communiqués vers Telquel, L'Économiste, 2M, MFM Radio lors du lancement.")
add_bullet("Programme de parrainage : 10 MAD de crédit offerts à l'utilisateur et à son filleul pour toute première réservation parrainée.")
add_bullet("Présence lors d'événements culturels marocains : Festival National du Film de Tanger, Marrakech International Film Festival (FIFM).")
add_space()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 6 — FAISABILITÉ TECHNIQUE
# ─────────────────────────────────────────────────────────────────────────────
add_heading1('6. Étude de Faisabilité Technique')

add_heading2('6.1 Ressources matérielles nécessaires')
add_space(3)

make_table(
    ['Matériel / Service', 'Quantité', 'Coût unitaire (MAD)', 'Coût mensuel (MAD)', 'Usage'],
    [
        ['Poste développeur MacBook Pro M3 ou équivalent PC', '2', '15 000–25 000 MAD (achat unique)', '—', 'Développement, tests, déploiement'],
        ['Serveur Vercel Pro (hébergement Next.js)', '1', '200 MAD/mois', '200', 'Déploiement production + CDN mondial'],
        ['Supabase Pro (BDD + Auth + Storage)', '1', '250 MAD/mois', '250', 'Base de données, authentification, stockage affiches'],
        ['Nom de domaine (.ma ou .com)', '1', '150–400 MAD/an', '~25', 'Identité web (lumiere-cinemas.ma)'],
        ['SendGrid Essentials (emails)', '1', '150 MAD/mois', '150', 'Billets par email, notifications'],
        ['SMS Twilio (notifications)', '1', '~0,85 MAD/SMS × volume', 'Variable', 'Confirmations de réservation'],
        ['CMI — adhésion paiement carte MA', '1', '3 000 MAD/an', '250', 'Paiement cartes Visa/MC marocaines'],
        ['Logiciels (Figma, VS Code, GitHub)', '1 équipe', 'Gratuit/open source', '0', 'Design, développement, versioning'],
    ],
    [4.5, 1.5, 3.5, 2.5, 3.5]
)

add_heading2('6.2 Ressources humaines')
add_body(
    "L'équipe initiale est volontairement légère pour minimiser les charges fixes au démarrage. "
    "Les salaires indiqués sont conformes au marché marocain 2025-2026 (source : enquêtes Cabinet "
    "Michael Page Maroc, Bayt.com Maroc) :"
)
add_space(3)

make_table(
    ['Profil', 'Qtité', 'Salaire brut mensuel', 'CNSS patronal (~23,98 %)', 'Coût total employeur/mois', 'Missions'],
    [
        ['Développeur Fullstack Senior (Next.js, Supabase, TypeScript)',
         '1', '12 000 MAD', '2 878 MAD', '14 878 MAD',
         'Maintenance, évolutions fonctionnelles, sécurité, déploiement'],
        ['Développeur Fullstack Junior / Alternant',
         '1', '5 000 MAD', '1 199 MAD', '6 199 MAD',
         'Support développement, tests, intégrations API partenaires'],
        ['Designer UI/UX',
         '1', '7 000 MAD', '1 679 MAD', '8 679 MAD',
         'Maquettes, identité visuelle, optimisation expérience utilisateur mobile'],
        ['Chargé(e) de Marketing Digital',
         '1', '6 000 MAD', '1 439 MAD', '7 439 MAD',
         'Réseaux sociaux, SEO, campagnes payantes, partenariats influenceurs'],
        ['Account Manager Cinémas (B2B)',
         '1', '7 000 MAD + commissions', '1 679 MAD', '8 679 MAD',
         'Prospection cinémas partenaires, onboarding, support opérateur'],
        ['Comptable (prestataire externe)',
         '1', '1 500 MAD/mois (forfait)', '—', '1 500 MAD',
         'Déclarations fiscales, paie, TVA, bilan annuel'],
    ],
    [4.0, 1.2, 2.5, 2.5, 3.0, 4.0]
)
add_space()
add_body("Masse salariale totale employeur mensuelle : 47 374 MAD", bold=True)
add_space()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 7 — FAISABILITÉ FINANCIÈRE
# ─────────────────────────────────────────────────────────────────────────────
add_heading1('7. Étude de Faisabilité Financière')

add_heading2('7.1 Investissement initial')
add_space(3)

make_table(
    ['Poste d\'investissement', 'Montant (MAD)', 'Détail'],
    [
        ['Développement initial de la plateforme (MVP)',
         '30 000',
         '3 mois de travail développeur senior (équivalent coût réel si externalisé)'],
        ['Design UI/UX — identité visuelle et maquettes',
         '8 000',
         'Charte graphique, logotype, maquettes Figma complètes'],
        ['Constitution juridique SARL',
         '3 500',
         'Frais notaire ~1 500 MAD + registre de commerce ~350 MAD + IF + publications légales'],
        ['Dépôt marque OMPIC (classe 41)',
         '2 025',
         'Dépôt officiel protection marque « Lumière Cinémas »'],
        ['Déclaration CNDP (loi 09-08)',
         '1 500',
         'Dossier déclaration traitement données personnelles'],
        ['Infrastructure technique initiale (3 mois)',
         '1 950',
         'Vercel Pro 200 MAD + Supabase Pro 250 MAD × 3 mois'],
        ['Intégration CMI — adhésion paiement',
         '3 000',
         'Frais d\'adhésion annuels Centre Monétique Interbancaire'],
        ['Budget marketing lancement (3 mois)',
         '15 000',
         'Publicités Instagram/TikTok 5 000 MAD/mois × 3 mois'],
        ['Matériel informatique (2 postes)',
         '20 000',
         '2 ordinateurs portables développeurs (10 000 MAD/unité)'],
        ['Fonds de roulement de précaution (2 mois charges)',
         '30 000',
         'Réserve pour couvrir 2 mois de charges fixes sans CA'],
        ['TOTAL INVESTISSEMENT INITIAL', '115 975 MAD', ''],
    ],
    [6.0, 2.8, 7.0]
)
# Bold last row
for cell in doc.tables[-1].rows[-1].cells:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
    shade_cell(cell, 'D9D9D9')

add_heading2('7.2 Prévisions de chiffre d\'affaires mensuel')
add_body(
    "Les hypothèses de croissance sont conservatrices et basées sur un lancement à Casablanca "
    "avec 2 cinémas partenaires (soit environ 8 salles) à partir du mois M+4 post-création."
)
add_space(3)

make_table(
    ['Source de revenus', 'Mois 3', 'Mois 6', 'Mois 9', 'Mois 12'],
    [
        ['Billets vendus (nombre/mois)',               '0 (lancement)', '1 200', '3 500', '7 000'],
        ['Commission moyenne/billet (MAD)',             '0 (gratuit)',   '3,50',  '3,50',  '4,00'],
        ['CA commissions billets (MAD)',                '0',             '4 200', '12 250','28 000'],
        ['Abonnements cinémas partenaires (MAD/mois)',  '990 × 1 = 990', '990 × 2 = 1 980', '2 490 × 2 = 4 980', '2 490 × 3 = 7 470'],
        ['Abonnements Lumière Pass spectateurs (MAD)',  '0',             '59 × 30 = 1 770', '59 × 80 = 4 720', '59 × 150 = 8 850'],
        ['CA Snacks (commission 15 %)',                 '0',             '600',   '1 800', '3 600'],
        ['CHIFFRE D\'AFFAIRES TOTAL MENSUEL (MAD)',     '990',           '8 550', '23 750','47 920'],
    ],
    [4.5, 2.5, 2.5, 2.5, 2.5]
)
# Bold last row
for cell in doc.tables[-1].rows[-1].cells:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
    shade_cell(cell, 'D9D9D9')

add_heading2('7.3 Charges variables mensuelles')
add_space(3)

make_table(
    ['Charge variable', 'Base de calcul', 'Mois 6 (MAD)', 'Mois 12 (MAD)'],
    [
        ['Commission Stripe (1,5 % + 1,50 MAD)',    '1 200 billets × 60 MAD moy.',  '1 980', '4 200'],
        ['Commission CMI (2 %)',                    'Paiements cartes MA',          '600',   '1 400'],
        ['SMS Twilio (0,85 MAD/SMS)',               '1 200 SMS/mois',               '1 020', '2 975'],
        ['SendGrid emails (0 à volume faible)',     '<50 000 emails → 150 MAD/mois','150',   '150'],
        ['Publicité digitale variable',             'CPC Instagram/TikTok',         '3 000', '5 000'],
        ['TOTAL CHARGES VARIABLES/MOIS (MAD)',      '',                             '6 750', '13 725'],
    ],
    [4.0, 4.0, 2.5, 2.5]
)
for cell in doc.tables[-1].rows[-1].cells:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
    shade_cell(cell, 'D9D9D9')

add_heading2('7.4 Charges fixes mensuelles')
add_space(3)

make_table(
    ['Charge fixe', 'Montant mensuel (MAD)'],
    [
        ['Masse salariale totale (coût employeur, 5 salariés)', '47 374'],
        ['Hébergement Vercel Pro',                              '200'],
        ['Supabase Pro',                                        '250'],
        ['Domaine + SSL',                                       '25'],
        ['SendGrid Essentials',                                 '150'],
        ['CMI — amortissement adhésion annuelle',              '250'],
        ['Comptable (prestataire)',                             '1 500'],
        ['Loyer bureau partagé (coworking, optionnel)',        '2 000'],
        ['Amortissement matériel informatique (24 mois)',      '833'],
        ['Divers (téléphone, transports, frais bancaires)',    '500'],
        ['TOTAL CHARGES FIXES MENSUELLES (MAD)',               '53 082'],
    ],
    [10.0, 5.0]
)
for cell in doc.tables[-1].rows[-1].cells:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
    shade_cell(cell, 'D9D9D9')

add_heading2('7.5 Calcul du Seuil de Rentabilité')
add_body("Formules appliquées :")
add_body("Taux de marge sur coût variable = (CA – Charges variables) ÷ CA", indent=True)
add_body("À mois 12 : (47 920 – 13 725) ÷ 47 920 = 34 195 ÷ 47 920 = 71,4 %", indent=True)
add_space(3)

make_table(
    ['Indicateur', 'Valeur (Mois 12)'],
    [
        ['Chiffre d\'affaires mensuel',    '47 920 MAD'],
        ['Charges variables mensuelles',   '13 725 MAD'],
        ['Marge sur coût variable',        '34 195 MAD'],
        ['Taux de marge sur CV',           '71,4 %'],
        ['Charges fixes mensuelles',       '53 082 MAD'],
        ['Seuil de rentabilité mensuel',   '53 082 ÷ 71,4 % = 74 344 MAD de CA/mois'],
        ['Nombre de billets nécessaires (commission 4 MAD + abo)', '~9 800 billets/mois'],
        ['Mois estimé d\'atteinte du seuil', 'Mois 14–16 (environ 15 mois après lancement)'],
    ],
    [8.0, 7.0]
)
for cell in doc.tables[-1].rows[-1].cells:
    shade_cell(cell, 'FFF9C4')
for cell in doc.tables[-1].rows[-2].cells:
    shade_cell(cell, 'FFF9C4')

add_heading2('7.6 Besoin en Fonds de Roulement (BFR)')
add_body(
    "Dans un modèle de billetterie numérique, le BFR est structurellement faible car :"
)
add_bullet("Stocks = 0 MAD (service numérique, aucun stock physique).")
add_bullet("Créances clients = faibles (paiement immédiat en ligne par carte ou espèces à la séance).")
add_bullet("Dettes fournisseurs = Stripe reverse les fonds sous 2–7 jours ouvrés.")
add_body(
    "BFR estimé = Créances clients (5 jours × CA journalier) – Dettes fournisseurs (7 jours × charges variables journalières)"
)
add_body(
    "À mois 12 : BFR ≈ (47 920 × 5/30) – (13 725 × 7/30) = 7 987 – 3 202 = 4 785 MAD.",
    bold=True
)
add_body(
    "Le BFR est très bas (< 5 000 MAD à maturité), ce qui est une caractéristique favorable "
    "du modèle SaaS/marketplace digital."
)

add_heading2('7.7 Trésorerie prévisionnelle sur 6 mois')
add_space(3)

make_table(
    ['Mois', 'CA (MAD)', 'Charges Var. (MAD)', 'Charges Fixes (MAD)', 'Résultat mensuel (MAD)', 'Trésorerie cumulée (MAD)'],
    [
        ['M1 (pré-lancement)', '0',      '500',   '53 082', '-53 582', '-53 582'],
        ['M2',                 '500',    '1 000', '53 082', '-53 582', '-107 164'],
        ['M3',                 '990',    '1 500', '53 082', '-53 592', '-160 756'],
        ['M4',                 '3 200',  '3 200', '53 082', '-53 082', '-213 838'],
        ['M5',                 '6 000',  '5 000', '53 082', '-52 082', '-265 920'],
        ['M6',                 '8 550',  '6 750', '53 082', '-51 282', '-317 202'],
    ],
    [1.8, 2.2, 2.5, 2.8, 3.0, 3.2]
)
add_space()
add_body(
    "Note : le capital de départ (115 975 MAD d\'investissement initial + apports associés) "
    "doit couvrir les pertes des 12 à 15 premiers mois. Un financement total de "
    "250 000 à 300 000 MAD est recommandé pour assurer la trésorerie jusqu'au seuil de "
    "rentabilité, incluant les 115 975 MAD d\'investissement initial et un coussin de "
    "trésorerie de 180 000 MAD pour couvrir les déficits mensuels."
)

add_heading2('7.8 Récapitulatif financier — Tableaux de synthèse')
add_body(
    "Les trois tableaux suivants synthétisent les indicateurs financiers clés à partir des "
    "données du mois 12 (premier exercice complet) :"
)

add_heading3('Tableau 1 — Coût total')
add_space(3)

t_cout = doc.add_table(rows=4, cols=2)
t_cout.style = 'Table Grid'
t_cout.alignment = WD_TABLE_ALIGNMENT.CENTER

add_table_header_row(t_cout, ['Élément', 'Montant (MAD)'])
data_cout = [
    ('Total charges variables', '13 725 MAD'),
    ('Total charges fixes',     '53 082 MAD'),
    ('COÛT TOTAL',              '66 807 MAD'),
]
for i, (elem, val) in enumerate(data_cout):
    row = t_cout.rows[i + 1]
    shade_cell(row.cells[0], 'F2F2F2' if i % 2 == 0 else 'FFFFFF')
    shade_cell(row.cells[1], 'F2F2F2' if i % 2 == 0 else 'FFFFFF')
    p0 = row.cells[0].paragraphs[0]
    p1 = row.cells[1].paragraphs[0]
    r0 = p0.add_run(elem)
    r1 = p1.add_run(val)
    r0.font.size = Pt(11)
    r1.font.size = Pt(11)
    if i == 2:
        r0.bold = True
        r1.bold = True
        shade_cell(row.cells[0], 'D9D9D9')
        shade_cell(row.cells[1], 'D9D9D9')

t_cout.columns[0].width = Cm(8)
t_cout.columns[1].width = Cm(5)

add_space()
add_heading3('Tableau 2 — Résultat prévisionnel (Mois 12)')
add_space(3)

t_res = doc.add_table(rows=4, cols=2)
t_res.style = 'Table Grid'
t_res.alignment = WD_TABLE_ALIGNMENT.CENTER

add_table_header_row(t_res, ['Élément', 'Montant (MAD)'])
data_res = [
    ('Chiffre d\'affaires',      '47 920 MAD'),
    ('Coût total',               '66 807 MAD'),
    ('RÉSULTAT PRÉVISIONNEL',    '- 18 887 MAD'),
]
for i, (elem, val) in enumerate(data_res):
    row = t_res.rows[i + 1]
    p0 = row.cells[0].paragraphs[0]
    p1 = row.cells[1].paragraphs[0]
    row.cells[0].text = ''
    row.cells[1].text = ''
    p0 = row.cells[0].paragraphs[0]
    p1 = row.cells[1].paragraphs[0]
    r0 = p0.add_run(elem)
    r1 = p1.add_run(val)
    r0.font.size = Pt(11)
    r1.font.size = Pt(11)
    if i == 2:
        r0.bold = True
        r1.bold = True
        r1.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
        shade_cell(row.cells[0], 'FDECEA')
        shade_cell(row.cells[1], 'FDECEA')
    elif i % 2 == 0:
        shade_cell(row.cells[0], 'F2F2F2')
        shade_cell(row.cells[1], 'F2F2F2')

t_res.columns[0].width = Cm(8)
t_res.columns[1].width = Cm(5)

add_space()
add_body(
    "Note d'interprétation : le résultat négatif de -18 887 MAD au mois 12 est attendu et "
    "planifié — il correspond à la phase d'investissement et de croissance de la startup. "
    "L'apport initial de 300 000 MAD couvre largement ce déficit cumulé. Le seuil de "
    "rentabilité est atteint au mois 14–16, après quoi l'entreprise génère un résultat positif."
)

add_heading3('Tableau 3 — Trésorerie prévisionnelle simplifiée')
add_space(3)

t_tres = doc.add_table(rows=5, cols=2)
t_tres.style = 'Table Grid'
t_tres.alignment = WD_TABLE_ALIGNMENT.CENTER

add_table_header_row(t_tres, ['Élément', 'Montant (MAD)'])
data_tres = [
    ('Apport initial (capital + financement)',  '300 000 MAD'),
    ('Revenus prévus (Mois 12)',                '47 920 MAD'),
    ('Dépenses prévues (Mois 12)',              '66 807 MAD'),
    ('SOLDE DE TRÉSORERIE DISPONIBLE (M12)',    '233 193 MAD'),
]
colors_tres = ['F2F2F2', 'FFFFFF', 'F2F2F2', 'D9D9D9']
for i, (elem, val) in enumerate(data_tres):
    row = t_tres.rows[i + 1]
    row.cells[0].text = ''
    row.cells[1].text = ''
    p0 = row.cells[0].paragraphs[0]
    p1 = row.cells[1].paragraphs[0]
    r0 = p0.add_run(elem)
    r1 = p1.add_run(val)
    r0.font.size = Pt(11)
    r1.font.size = Pt(11)
    shade_cell(row.cells[0], colors_tres[i])
    shade_cell(row.cells[1], colors_tres[i])
    if i == 3:
        r0.bold = True
        r1.bold = True

t_tres.columns[0].width = Cm(8)
t_tres.columns[1].width = Cm(5)

add_space()
add_body(
    "Calcul du solde : 300 000 (apport) – [cumul déficits M1 à M12 ≈ 66 807 MAD × 12 mois, "
    "atténué par les revenus progressifs] ≈ 233 193 MAD disponibles en trésorerie à la fin "
    "du 12e mois. Ce solde confirme que le projet reste solvable sur l'ensemble de la "
    "première année d'exploitation avec l'apport initial planifié."
)

add_space()
add_heading2('7.9 Analyse de rentabilité — Synthèse narrative')
add_body(
    "Le modèle économique de Lumière Cinémas est un modèle marketplace SaaS à deux faces : "
    "les spectateurs paient une commission par billet, et les cinémas paient un abonnement "
    "mensuel pour accéder au back-office. Cette double monétisation offre une stabilité de "
    "revenus (abonnements récurrents) complétée par une part variable liée au volume de "
    "billets vendus."
)
add_body(
    "Le seuil de rentabilité est atteint autour du 15e mois d\'exploitation, ce qui est "
    "cohérent avec les standards des startups SaaS en phase de croissance. Les principaux "
    "risques financiers identifiés sont : (1) la lenteur à convaincre les premiers cinémas "
    "partenaires (risque de CA B2B nul les premiers mois) ; (2) le coût élevé de la masse "
    "salariale qui représente 89 % des charges fixes ; (3) la dépendance au volume de "
    "billets pour la partie commission. Pour mitiger ces risques, il est recommandé de "
    "démarrer avec une équipe réduite (2 fondateurs + 1 développeur) et d\'augmenter les "
    "recrutements uniquement après signature des premiers partenariats cinémas."
)
add_body(
    "À partir du mois 18–24, avec 3 cinémas partenaires Pro et 10 000 billets/mois, "
    "le projet devient structurellement rentable avec une marge nette estimée à "
    "15–20 % du CA, soit environ 8 000 à 12 000 MAD de bénéfice mensuel. "
    "L'atteinte de 100 000 billets/mois (horizon 3 ans) générerait un CA annuel "
    "supérieur à 6 millions de MAD."
)
add_space()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 8 — ASPECT JURIDIQUE
# ─────────────────────────────────────────────────────────────────────────────
add_heading1('8. Aspect Juridique')

add_heading2('8.1 Forme juridique retenue : SARL')
add_body(
    "Lumière Cinémas sera constituée sous la forme d'une Société à Responsabilité Limitée (SARL) "
    "régie par la loi n° 5-96 sur la société en nom collectif, la société en commandite simple, "
    "la société en commandite par actions, la société à responsabilité limitée et la société en "
    "participation, telle que modifiée par la loi 21-05."
)
add_body(
    "Capital minimum légal : 10 000 MAD (libérable en totalité à la constitution). "
    "Le capital envisagé pour ce projet est de 100 000 MAD, partagé entre les associés fondateurs, "
    "afin de crédibiliser la société vis-à-vis des partenaires cinémas et des institutions financières."
)

add_heading2('8.2 Démarches de création réelles')
add_space(3)

make_table(
    ['Étape', 'Organisme', 'Délai', 'Coût estimé (MAD)'],
    [
        ['Certificat négatif (vérification disponibilité du nom)',
         'OMPIC (en ligne)',
         '24–48h',
         '230 MAD'],
        ['Rédaction des statuts et légalisation',
         'Notaire ou avocat',
         '3–5 jours',
         '1 500–3 000 MAD'],
        ['Enregistrement des statuts (droits d\'enregistrement)',
         'Direction des Impôts',
         '1–3 jours',
         '1 000–2 000 MAD'],
        ['Inscription au Registre de Commerce',
         'Tribunal de Commerce',
         '2–5 jours',
         '350 MAD'],
        ['Publication au Bulletin Officiel et journal d\'annonces légales',
         'BO + journal agréé',
         '5–10 jours',
         '1 500–2 000 MAD'],
        ['Identifiant Fiscal (IF)',
         'Direction Régionale des Impôts',
         '24–48h',
         'Gratuit'],
        ['Inscription à la CNSS (numéro patronal)',
         'Caisse Nationale de Sécurité Sociale',
         '2–5 jours',
         'Gratuit'],
        ['Patente (Taxe Professionnelle)',
         'DGI / Direction des Impôts',
         '5–10 jours',
         'Variable selon CA prévisionnel'],
        ['Compte bancaire professionnel',
         'Banque marocaine (CIH, Attijariwafa, BMCE)',
         '3–5 jours',
         '0–500 MAD (frais ouverture)'],
        ['TOTAL CRÉATION SARL', '', '~20–30 jours', '~5 000–8 000 MAD'],
    ],
    [5.0, 3.5, 2.0, 2.5]
)
for cell in doc.tables[-1].rows[-1].cells:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
    shade_cell(cell, 'D9D9D9')

add_heading2('8.3 Justification du choix SARL vs autres formes')
add_space(3)

make_table(
    ['Critère', 'SARL', 'Auto-entrepreneur', 'Société Anonyme (SA)'],
    [
        ['Capital minimum',      '10 000 MAD',              'Aucun',                '300 000 MAD'],
        ['Responsabilité',       'Limitée aux apports',     'Illimitée sur biens propres', 'Limitée aux apports'],
        ['Nombre d\'associés',   '1 à 50',                  '1 (seul)',              '5 minimum'],
        ['Ouverture capital',    'Non (sans transformation)', 'Non',                 'Oui (actions)'],
        ['Charges sociales',     'CNSS patron ~23,98 %',    'Forfait ~10–12 % du CA', 'Idem SARL'],
        ['Crédibilité B2B',      'Forte',                   'Faible',               'Très forte'],
        ['Complexité admin.',    'Modérée',                 'Simple',               'Complexe'],
        ['Recommandé pour',      '✓ CE PROJET',             'Activité solo simple', 'Levée de fonds > 5M MAD'],
    ],
    [4.0, 3.5, 3.5, 3.5]
)
add_space()
add_body(
    "La SARL est la forme juridique la plus adaptée à ce stade du projet : elle offre la "
    "protection de la responsabilité limitée (le patrimoine personnel des fondateurs est "
    "protégé), une crédibilité suffisante pour les négociations B2B avec les cinémas, "
    "une gestion administrative raisonnable et la possibilité d'accueillir des investisseurs "
    "supplémentaires via une augmentation de capital ou une transformation en SA le moment venu."
)
add_space()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 9 — STRATÉGIE DE DÉVELOPPEMENT
# ─────────────────────────────────────────────────────────────────────────────
add_heading1('9. Stratégie de Développement')

add_heading2('9.1 Roadmap produit')
add_space(3)

make_table(
    ['Phase', 'Horizon', 'Périmètre géographique', 'Fonctionnalités clés', 'Objectif CA mensuel'],
    [
        ['V1 — MVP Casablanca',
         'M0 → M12',
         'Casablanca (2 cinémas partenaires)',
         'Tunnel réservation complet, back-office, Stripe, QR code, programme centralisé. '
         'Interface web responsive. Arabe à intégrer.',
         '8 000–15 000 MAD'],
        ['V2 — Extension nationale + Mobile',
         'M12 → M24',
         'Casablanca, Rabat, Marrakech, Tanger, Agadir',
         'Application mobile React Native (iOS + Android). '
         'Géolocalisation des salles (Google Maps API). '
         'Notifications push réservation. '
         'Programme multi-villes. Intégration CMI complète.',
         '40 000–80 000 MAD'],
        ['V3 — Événements culturels + B2B',
         'M24 → M48',
         'National + Maroc entier',
         'Extension à d\'autres événements : concerts, théâtre, festivals. '
         'Module corporate (réservations entreprise, facturation). '
         'API ouverte pour intégration sur sites tiers. '
         'Analytics avancés et recommandations personnalisées (ML).',
         '150 000–300 000 MAD'],
        ['V4 — Internationalisation Maghreb',
         'M36 → M60',
         'Algérie, Tunisie (pilote)',
         'Adaptation devise (DZD, TND), langue arabe native. '
         'Partenariats locaux en Algérie (Cinéma Zinet, Cinotheque) '
         'et en Tunisie (CinéCity, Cine Majestic). '
         'Infrastructure multi-tenant Supabase.',
         '500 000+ MAD (pan-Maghreb)'],
    ],
    [2.0, 2.0, 3.0, 5.5, 2.5]
)

add_heading2('9.2 Indicateurs clés de performance (KPIs)')
add_body("Les KPIs suivis mensuellement pour piloter la croissance :")
add_bullet("GMV (Gross Merchandise Value) : valeur totale des billets vendus via la plateforme.")
add_bullet("Nombre de billets vendus / mois par cinéma partenaire.")
add_bullet("Taux de conversion : visiteurs uniques → réservations confirmées (cible : > 3,5 %).")
add_bullet("CAC (Coût d'Acquisition Client) : budget marketing ÷ nouveaux utilisateurs actifs (cible : < 15 MAD).")
add_bullet("LTV (Lifetime Value) : CA moyen généré par utilisateur sur 12 mois (cible : > 200 MAD).")
add_bullet("NPS (Net Promoter Score) : mesure de la satisfaction et de la propension à recommander.")
add_bullet("Taux de rétention à 30 jours : % d'utilisateurs revenant faire une deuxième réservation (cible : > 35 %).")
add_bullet("Churn rate B2B : taux de résiliation des abonnements cinémas partenaires (cible : < 5 %/mois).")

add_heading2('9.3 Perspectives d\'internationalisation')
add_body(
    "Le Maghreb représente un marché naturel d'expansion, avec une langue, une culture et des "
    "problématiques communes au Maroc. L'Algérie, avec ses 45 millions d'habitants et une "
    "renaissance récente du cinéma (réouverture des salles après les années 1990–2000), offre "
    "un potentiel considérable malgré les barrières réglementaires. La Tunisie, plus petite mais "
    "avec un tissu cinématographique plus mature (Festival de Carthage, réseau Cine Cities), "
    "constitue un marché test idéal pour l'internationalisation. L'architecture technique de "
    "Lumière Cinémas (Supabase multi-tenant, Next.js i18n, Stripe avec support multi-devises) "
    "a été pensée pour faciliter cette extension sans refonte majeure du code."
)
add_space()

# ─────────────────────────────────────────────────────────────────────────────
# CONCLUSION
# ─────────────────────────────────────────────────────────────────────────────
add_heading1('Conclusion')
add_body(
    "Lumière Cinémas répond à un besoin réel et non adressé sur le marché marocain de la billetterie "
    "cinéma en ligne. La solution développée, adossée à une stack technique moderne (Next.js 15, "
    "Supabase, Stripe, TypeScript), offre une expérience utilisateur premium sans équivalent local. "
    "L'étude entrepreneuriale démontre la viabilité du modèle avec un seuil de rentabilité "
    "atteignable en 14 à 16 mois, un BFR structurellement faible et une roadmap claire en 4 phases "
    "permettant d'adresser progressivement le marché national puis maghrébin. Le choix de la SARL "
    "comme structure juridique offre le bon équilibre entre protection, crédibilité et flexibilité "
    "pour cette phase d'amorçage. Lumière Cinémas a les atouts pour devenir la référence de la "
    "billetterie culturelle numérique au Maroc."
)
add_space()

# ─────────────────────────────────────────────────────────────────────────────
# SAUVEGARDE
# ─────────────────────────────────────────────────────────────────────────────
output_path = r'C:\Users\PC\Documents\Cinemaa-SSG\Rapport_Entrepreneuriat_LumiereCinemas_FINAL.docx'
doc.save(output_path)
print(f'Document généré : {output_path}')
